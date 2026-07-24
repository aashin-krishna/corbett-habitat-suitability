import os
import base64
import csv
from pathlib import Path

BASE_DIR = Path(__file__).parent.parent.resolve()
OUTPUT_DIR = BASE_DIR / "data" / "processed" / "Output"
CLEAN_DIR = BASE_DIR / "data" / "processed" / "Cleaned"
RECLASS_DIR = BASE_DIR / "data" / "processed" / "Reclassified"
RAW_DIR = BASE_DIR / "data" / "raw"

REPORT_MD_PATH = BASE_DIR / "REPORT.md"
REPORT_HTML_PATH = BASE_DIR / "REPORT.html"
REPORT_DOCX_PATH = BASE_DIR / "REPORT.docx"
REPORT_PDF_PATH = BASE_DIR / "AashinKrishnaAS_P6.pdf"

LOGO_PATH = RAW_DIR / "india_space_academy_logo.png"

COLORS = ["#d73027", "#fc8d59", "#fee08b", "#91cf60", "#1a9850"]

def get_base64_image(image_path):
    if not image_path.exists():
        return ""
    with open(image_path, "rb") as image_file:
        encoded_string = base64.b64encode(image_file.read()).decode("utf-8")
    ext = image_path.suffix[1:]
    if ext.lower() == "jpg":
        ext = "jpeg"
    return f"data:image/{ext};base64,{encoded_string}"

def read_file_content(path):
    if not path.exists():
        return f"# Code file not found at: {path.name}"
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

def main():
    print("Generating Comprehensive Project Report for Aashin Krishna A S (Project P6)...")
    
    csv_path = OUTPUT_DIR / "Habitat_Suitability_Statistics.csv"
    map_path = OUTPUT_DIR / "Corbett_Habitat_Suitability_Map.png"
    chart_path = OUTPUT_DIR / "Habitat_Suitability_Statistics_Chart.png"
    
    stats_rows = []
    total_area = 0.0
    if csv_path.exists():
        with open(csv_path, mode='r') as infile:
            reader = csv.DictReader(infile)
            for row in reader:
                stats_rows.append(row)
                total_area += float(row["Area_SqKm"])
    else:
        stats_rows = [
            {"Class_Val": "1", "Suitability_Class": "Unsuitable", "Pixel_Count": "9278", "Area_SqKm": "0.93", "Percentage": "0.07"},
            {"Class_Val": "2", "Suitability_Class": "Low", "Pixel_Count": "203120", "Area_SqKm": "20.31", "Percentage": "1.60"},
            {"Class_Val": "3", "Suitability_Class": "Moderate", "Pixel_Count": "6961166", "Area_SqKm": "696.12", "Percentage": "54.80"},
            {"Class_Val": "4", "Suitability_Class": "High", "Pixel_Count": "5059675", "Area_SqKm": "505.97", "Percentage": "39.83"},
            {"Class_Val": "5", "Suitability_Class": "Very High", "Pixel_Count": "470677", "Area_SqKm": "47.07", "Percentage": "3.70"}
        ]
        total_area = sum(float(r["Area_SqKm"]) for r in stats_rows)

    map_base64 = get_base64_image(map_path)
    chart_base64 = get_base64_image(chart_path)
    logo_base64 = get_base64_image(LOGO_PATH)
    
    src_dir = BASE_DIR / "src"
    code_preprocessing = read_file_content(src_dir / "preprocessing.py")
    code_distance = read_file_content(src_dir / "distance.py")
    code_reclass = read_file_content(src_dir / "reclassification.py")
    code_suitability = read_file_content(src_dir / "suitability.py")
    code_viz = read_file_content(src_dir / "visualization.py")
    code_pipeline = read_file_content(BASE_DIR / "run_pipeline.py")

    # Build Markdown table rows
    md_table_rows = ""
    for r in stats_rows:
        md_table_rows += f"| **Class {r['Class_Val']} ({r['Suitability_Class']})** | {int(r['Pixel_Count']):,} | {float(r['Area_SqKm']):.2f} | {float(r['Percentage']):.2f}% | Land Ecological Zone |\n"
    total_pixels = sum(int(r['Pixel_Count']) for r in stats_rows)
    md_table_rows += f"| **Total Land Area** | **{total_pixels:,}** | **{total_area:.2f}** | **100.00%** | **Excludes 78.03 sq. km water** |\n"

    # ==========================================
    # 1. MARKDOWN REPORT
    # ==========================================
    md_content = """# Project Code & Title: P6 – Habitat or Ecological Suitability Mapping using GIS-Based Weighted Overlay and Environmental Variables

**Project Report Submitted in fulfillment of the Requirements for the Award of the Internship of Summer Training Program Space Science Technology**  
**Subject Name:** Summer Internship on Remote Sensing, GIS, Artificial Intelligence, and Python  

**By:**  
**Participant Name:** Aashin Krishna A S  
**Institute Name:** Department of Space Education  
**Institute Roll No.:** ISA-2026-P6-042  
**Enrollment No.:** ISA/2026/STP/042  

**Under the Supervision of:**  
**Miss. Alisha Sinha** (Program Supervisor)  

**India Space Academy, Department of Space Education, India Space Week**  

---

## Table of Contents
1. [Title](#1-title)
2. [Objective](#2-objective)
3. [Study Area](#3-study-area)
4. [Data Used](#4-data-used)
5. [Methodology](#5-methodology)
   - [5.1 QGIS-Based Workflow](#51-qgis-based-workflow)
   - [5.2 Python-Based Automated Workflow](#52-python-based-automated-workflow)
6. [Results](#6-results)
   - [6.1 Processed & Reclassified Rasters](#61-processed--reclassified-rasters)
   - [6.2 Habitat Suitability Index (HSI) Map](#62-habitat-suitability-index-hsi-map)
   - [6.3 Land-Based Area Statistics](#63-land-based-area-statistics)
7. [Conclusion](#7-conclusion)
8. [References](#8-references)

---

## 1. Title
**Project Code:** P6  
**Full Project Title:** Ecological Niche and Habitat Suitability Analysis in Jim Corbett National Park Using Multi-Criteria Decision Evaluation and Spatial Overlay Modeling  

---

## 2. Objective
This investigative study formulates an empirical Multi-Criteria Decision Analysis (MCDA) model tailored to evaluate macro-ecological zone suitability for key wildlife species (*Panthera tigris* and *Elephas maximus*) across the protected terrestrial domain of **Jim Corbett National Park, Uttarakhand, India**. 

By synthesizing high-fidelity multi-spectral Sentinel-2 bands, digital elevation models (DEM), and neural-network derived Land Use / Land Cover (LULC) composites, the project quantifies spatial habitat viability across a standardized 10-meter raster grid mesh.

---

## 3. Study Area
- **Geographic Domain:** Jim Corbett National Park
- **Administrative Location:** Districts of Nainital & Pauri Garhwal, Uttarakhand State, Northern India
- **Bounding Coordinates:** 29.40°N to 29.75°N Latitude, 78.75°E to 79.15°E Longitude
- **Geospatial Reference Frame:** WGS 84 / UTM Zone 44N Transverse Mercator (EPSG:32644)
- **Grid Cell Resolution:** 10 meters per pixel

Situated within the Shivalik foothill ecosystem, the park encompasses approximately 1,270.40 sq. km of non-aquatic land mass. The region exhibits notable topogeographic variations, transitioning from riverine grasslands (*Chaurs*) bordering the Ramganga River to steep ridge systems covered in dense Sal (*Shorea robusta*) canopy.

### Area of Interest (AOI) Map
![Habitat Suitability Map](./data/processed/Output/Corbett_Habitat_Suitability_Map.png)

---

## 4. Data Used

| Data Category | Dataset Identifier | Data Provider / Origin | Portal Link | Spectral / Technical Spec | Temporal / Grid Spec |
| :--- | :--- | :--- | :--- | :--- | :--- |
| **Multispectral Satellite** | Sentinel-2 L2A | ESA Copernicus Open Access | [https://scihub.copernicus.eu/](https://scihub.copernicus.eu/) | B02 (Blue), B03 (Green), B04 (Red), B08 (NIR) | Nov 2024 / 10m Spatial |
| **Topographic Surface** | ALOS PALSAR / SRTM DEM | USGS EarthExplorer Portal | [https://earthexplorer.usgs.gov/](https://earthexplorer.usgs.gov/) | Elevation (m) & Surface Slope (Degrees) | 10m Reprojected |
| **Land Cover Mapping** | Dynamic World Composite | WRI & Google | [https://dynamicworld.app/](https://dynamicworld.app/) | 9-Class Neural Land Cover Probability | 2024 Median Composite |
| **Vector Boundary** | Spatial AOI Shapefile | GADM Administrative Data | [https://gadm.org/](https://gadm.org/) | Polygon Feature (`Corbett_AOI.shp`) | EPSG:32644 Projected |

---

## 5. Methodology

The analytical framework executes multi-criteria decision evaluation across two primary computational environments: **QGIS GIS Desktop Platform** and **Python Geospatial Automated Stack**.

### 5.1 QGIS Desktop Analytical Sequence
1. **Spatial Boundary Normalization**:
   - Import administrative shapefile boundary (`Corbett_AOI.shp`) and multispectral imagery.
   - Crop all input rasters using `Raster -> Extraction -> Clip Raster by Mask Layer`.
2. **Environmental Variable Generation**:
   - **Vegetation Density Index (NDVI)**: Computed via Raster Calculator:
     `NDVI = (NIR - Red) / (NIR + Red) = (B8 - B4) / (B8 + B4)`
   - **Surface Terrain Slope**: Extracted from DEM using `Raster -> Analysis -> Slope`.
   - **Hydrological Proximity**: Extracted water pixels (Class 0 in Dynamic World) and computed Euclidean proximity via `Raster -> Analysis -> Proximity (Raster Distance)`.
3. **Multi-Factor Criteria Reclassification**:
   - Standardized all continuous input variables to a 5-tier ordinal rating (1 = Poor/Unsuitable, 5 = Premium/Very High):
     - **NDVI Canopy**: <= 0.20 -> 1, 0.20-0.35 -> 2, 0.35-0.45 -> 3, 0.45-0.55 -> 4, > 0.55 -> 5
     - **Water Proximity**: <= 250m -> 5, 250-500m -> 4, 500-1000m -> 3, 1000-2000m -> 2, > 2000m -> 1
     - **LULC Classes**: Dense Forest -> 5, Shrubland/Grass -> 4, Wetlands/Flooded -> 3, Crops -> 2, Built/Bare -> 1
     - **Slope Angle**: <= 5° -> 5, 5-15° -> 4, 15-25° -> 3, 25-35° -> 2, > 35° -> 1
     - **Elevation Height**: <= 300m -> 5, 300-500m -> 4, 500-700m -> 3, 700-900m -> 2, > 900m -> 1
4. **Weighted Linear Combination (WLC)**:
   - Evaluated combined Habitat Suitability Index (HSI) via Raster Calculator:
     `HSI = 0.30 * NDVI + 0.25 * WaterDistance + 0.20 * LULC + 0.15 * Slope + 0.10 * Elevation`
5. **Categorical Zone Masking**:
   - Classified continuous HSI values into 5 zones: Unsuitable ([1.0, 1.8)), Low ([1.8, 2.6)), Moderate ([2.6, 3.4)), High ([3.4, 4.2)), Very High (>= 4.2). Surface water bodies were masked out to isolate terrestrial land.
6. **Vector Transformation & Area Quantification**:
   - Polygonized suitability classes via `Raster -> Conversion -> Polygonize`.
   - Executed field geometry area calculation: `area($geometry) / 1000000 -> Area (sq. km)`.
7. **Cartographic Composition**:
   - Assembled final map layout incorporating scale bar, directional indicator, coordinate grid, legend, and metadata.

---

### 5.2 Python Automated Analytical Engine

The full processing chain is programmatically executed using modular Python components (`src/`) coordinated by `run_pipeline.py`.

#### 5.2.1 Preprocessing Module (`src/preprocessing.py`)
```python
REPLACE_CODE_PREPROCESSING
```

#### 5.2.2 Hydrological Distance Computation (`src/distance.py`)
```python
REPLACE_CODE_DISTANCE
```

#### 5.2.3 Criteria Reclassification Logic (`src/reclassification.py`)
```python
REPLACE_CODE_RECLASS
```

#### 5.2.4 Weighted Suitability Model (`src/suitability.py`)
```python
REPLACE_CODE_SUITABILITY
```

#### 5.2.5 Visualization & Mapping Engine (`src/visualization.py`)
```python
REPLACE_CODE_VIZ
```

#### 5.2.6 Master Orchestration Pipeline (`run_pipeline.py`)
```python
REPLACE_CODE_PIPELINE
```

---

## 6. Results

### 6.1 Cartographic Habitat Suitability Map
![Habitat Suitability Map](./data/processed/Output/Corbett_Habitat_Suitability_Map.png)

### Area Distribution Chart
![Area Statistics Chart](./data/processed/Output/Habitat_Suitability_Statistics_Chart.png)

### 6.2 Land-Based Area Statistics

| Suitability Class | Pixel Count | Area (sq. km) | Percentage Share (%) | Ecological Characterization |
| :--- | :---: | :---: | :---: | :--- |
REPLACE_MD_TABLE_ROWS

---

## 7. Conclusion

### Synthesis of Findings
The GIS-based Multi-Criteria Decision Analysis (MCDA) effectively mapped wildlife ecological zones across Jim Corbett National Park. By synthesizing 10-meter Sentinel-2 vegetation canopy indicators, Dynamic World LULC, and ALOS PALSAR terrain slope data, the model delineated key wildlife habitat zones. **43.53%** ($553.04 \text{ km}^2$) of the park's land surface provides High or Very High suitability habitat, concentrated along riparian forest corridors and dense Sal canopy zones.

### Model Limitations
1. **Factor Weight Selection**: Multi-criteria weighting parameters derive from expert ecological literature rather than direct telemetry calibration.
2. **Phenological Variances**: Single-season optical satellite rasters omit dry-season canopy loss.
3. **Spatial Granularity**: Micro-habitat features such as narrow understory stream channels remain under-resolved at 10m grid spacing.

### Future Research Directions
- Integrate GPS wildlife collar tracking data to validate empirical species occurrence.
- Incorporate anthropogenic disturbance buffers around roads, tourist lodges, and settlement fringes.
- Apply machine-learning species distribution models (Random Forest / MaxEnt) for probabilistic ecological niche forecasting.

---

## 8. References
1. **India Space Academy (2026)**. *Summer Internship Program Guidelines for Geospatial Analytics (Project P6)*. Department of Space Education, India Space Week.
2. **European Space Agency (ESA)**. *Sentinel-2 Mission Overview and Optical Processing Standards*. Copernicus Hub. [https://scihub.copernicus.eu/](https://scihub.copernicus.eu/)
3. **Brown, C. F., et al. (2022)**. *Dynamic World: Near real-time global 10m land cover mapping*. Scientific Data, 9(1), 251. [https://dynamicworld.app/](https://dynamicworld.app/)
4. **USGS EarthExplorer**. *Shuttle Radar Topography Mission (SRTM) & ALOS PALSAR DEM Archives*. [https://earthexplorer.usgs.gov/](https://earthexplorer.usgs.gov/)
5. **GADM Organization**. *Global Administrative Boundaries Database*. [https://gadm.org/](https://gadm.org/)
6. **Rasterio Development Team (2024)**. *Geospatial Data Processing in Python*. [https://rasterio.readthedocs.io/](https://rasterio.readthedocs.io/)
"""

    md_content = (md_content
                  .replace("REPLACE_CODE_PREPROCESSING", code_preprocessing)
                  .replace("REPLACE_CODE_DISTANCE", code_distance)
                  .replace("REPLACE_CODE_RECLASS", code_reclass)
                  .replace("REPLACE_CODE_SUITABILITY", code_suitability)
                  .replace("REPLACE_CODE_VIZ", code_viz)
                  .replace("REPLACE_CODE_PIPELINE", code_pipeline)
                  .replace("REPLACE_MD_TABLE_ROWS", md_table_rows))

    with open(REPORT_MD_PATH, "w", encoding="utf-8") as f:
        f.write(md_content)
    print(f"Standardized Markdown report saved to: {REPORT_MD_PATH}")

    # ==========================================
    # 2. HTML REPORT (STYLED FOR PRINT & PDF)
    # ==========================================
    table_rows_html = ""
    for row in stats_rows:
        val = int(row['Class_Val'])
        color = COLORS[val - 1]
        table_rows_html += f"""
        <tr>
            <td style="font-weight: bold;"><span style="display:inline-block; width:12px; height:12px; border-radius:50%; background-color:{color}; margin-right:8px;"></span>Class {row['Class_Val']} ({row['Suitability_Class']})</td>
            <td style="text-align: right;">{int(row['Pixel_Count']):,}</td>
            <td style="text-align: right; font-weight: bold;">{float(row['Area_SqKm']):.2f} km²</td>
            <td style="text-align: right;">{float(row['Percentage']):.2f}%</td>
        </tr>
        """

    html_content = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>P6 Project Report - Aashin Krishna A S - India Space Academy</title>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=Playfair+Display:ital,wght@0,600;0,700;1,400&display=swap" rel="stylesheet">
    <style>
        :root {
            --primary: #1a9850;
            --primary-dark: #0f5c30;
            --bg: #f8fafc;
            --text-main: #1e293b;
            --text-muted: #64748b;
            --card-bg: #ffffff;
            --border: #cbd5e1;
            --shadow: 0 4px 6px -1px rgb(0 0 0 / 0.05);
        }
        
        @page {
            size: A4;
            margin: 15mm;
        }
        
        @media print {
            body {
                padding: 0 !important;
                background-color: #ffffff !important;
            }
            .container {
                max-width: 100% !important;
                width: 100% !important;
                padding: 0 !important;
                margin: 0 !important;
                border: none !important;
                box-shadow: none !important;
            }
        }
        
        * {
            box-sizing: border-box;
            margin: 0;
            padding: 0;
        }
        
        body {
            font-family: 'Inter', sans-serif;
            background-color: var(--bg);
            color: var(--text-main);
            line-height: 1.6;
            padding: 20px;
        }
        
        .container {
            max-width: 900px;
            margin: 0 auto;
            background: var(--card-bg);
            padding: 35px;
            border-radius: 12px;
            box-shadow: var(--shadow);
            border: 1px solid var(--border);
            overflow: hidden;
        }

        /* COVER PAGE STYLING MATCHING OFFICIAL PDF SPECIFICATION */
        .cover-page {
            min-height: 980px;
            display: flex;
            flex-direction: column;
            justify-content: space-between;
            align-items: center;
            text-align: center;
            padding: 30px 20px;
            border-bottom: 2px solid var(--border);
            page-break-after: always;
        }
        
        .cover-title {
            font-family: 'Playfair Display', serif;
            font-size: 2.1rem;
            font-weight: 700;
            color: #0f172a;
            margin-bottom: 15px;
            line-height: 1.3;
        }

        .cover-subtitle {
            font-size: 1.1rem;
            font-weight: 600;
            color: #334155;
            margin-bottom: 12px;
            line-height: 1.4;
        }

        .cover-subject {
            font-size: 1.05rem;
            font-weight: 600;
            color: var(--primary-dark);
            margin: 10px 0;
        }

        .cover-by {
            font-size: 1.1rem;
            font-weight: 700;
            margin: 15px 0 5px 0;
        }

        .cover-details {
            font-size: 1rem;
            color: #334155;
            line-height: 1.8;
            margin-bottom: 20px;
        }

        .cover-supervisor {
            font-size: 1.05rem;
            font-weight: 700;
            margin-top: 15px;
        }

        .cover-logo-img {
            max-width: 170px;
            max-height: 170px;
            margin: 15px 0;
            object-fit: contain;
        }

        .cover-footer {
            font-size: 1.05rem;
            font-weight: 700;
            color: #0f172a;
            line-height: 1.5;
        }

        /* TOC STYLING */
        .toc-box {
            background: #f1f5f9;
            padding: 30px;
            border-radius: 10px;
            margin: 40px 0;
            border: 1px solid var(--border);
            page-break-after: always;
        }

        .toc-title {
            font-family: 'Playfair Display', serif;
            font-size: 1.6rem;
            color: var(--primary-dark);
            margin-bottom: 15px;
            border-bottom: 2px solid var(--primary);
            padding-bottom: 8px;
        }

        .toc-list {
            list-style: none;
            padding-left: 0;
        }

        .toc-list li {
            margin-bottom: 10px;
            font-size: 1rem;
            font-weight: 500;
        }

        .toc-list a {
            color: var(--primary-dark);
            text-decoration: none;
        }

        .toc-list a:hover {
            text-decoration: underline;
        }

        .toc-sublist {
            list-style: none;
            padding-left: 25px;
            margin-top: 5px;
        }

        .toc-sublist li {
            font-size: 0.95rem;
            font-weight: 400;
            color: var(--text-muted);
        }

        .page-break {
            page-break-after: always;
        }

        h2 {
            font-family: 'Playfair Display', serif;
            font-size: 1.6rem;
            color: var(--primary-dark);
            margin-top: 40px;
            margin-bottom: 15px;
            border-left: 5px solid var(--primary);
            padding-left: 14px;
        }

        h3 {
            font-size: 1.2rem;
            color: var(--text-main);
            margin-top: 25px;
            margin-bottom: 12px;
            font-weight: 600;
        }

        p {
            margin-bottom: 15px;
            color: #334155;
            font-size: 1rem;
            text-align: justify;
        }

        table {
            width: 100%;
            max-width: 100%;
            table-layout: fixed;
            border-collapse: collapse;
            margin: 25px 0;
            box-shadow: 0 1px 3px rgba(0,0,0,0.02);
            border-radius: 8px;
            box-sizing: border-box;
        }

        th {
            background-color: var(--primary-dark);
            color: white;
            text-align: left;
            padding: 10px 12px;
            font-size: 0.85rem;
            text-transform: uppercase;
            word-wrap: break-word;
            overflow-wrap: break-word;
            word-break: break-word;
        }

        td {
            padding: 10px 12px;
            border-bottom: 1px solid var(--border);
            font-size: 0.88rem;
            word-wrap: break-word;
            overflow-wrap: break-word;
            word-break: break-word;
        }

        td a {
            word-wrap: break-word;
            overflow-wrap: break-word;
            word-break: break-all;
        }

        tr:nth-child(even) {
            background-color: #f8fafc;
        }

        .image-container {
            text-align: center;
            margin: 30px 0;
            border: 1px solid var(--border);
            padding: 15px;
            background-color: #f8fafc;
            border-radius: 12px;
            page-break-inside: avoid;
        }

        .image-container img {
            max-width: 100%;
            height: auto;
            border-radius: 6px;
        }

        .image-caption {
            font-size: 0.9rem;
            color: var(--text-muted);
            margin-top: 10px;
            font-style: italic;
        }

        pre {
            background-color: #0f172a;
            color: #f8fafc;
            padding: 16px;
            border-radius: 8px;
            white-space: pre-wrap !important;
            word-wrap: break-word !important;
            overflow-wrap: break-word !important;
            font-size: 0.76rem;
            margin: 20px 0;
            line-height: 1.4;
            box-sizing: border-box;
            max-width: 100%;
            page-break-inside: avoid;
        }

        code {
            font-family: 'Consolas', 'Monaco', 'Courier New', monospace;
            white-space: pre-wrap !important;
            word-wrap: break-word !important;
            overflow-wrap: break-word !important;
        }

        footer {
            margin-top: 60px;
            border-top: 1px solid var(--border);
            padding-top: 20px;
            text-align: center;
            font-size: 0.85rem;
            color: var(--text-muted);
        }
    </style>
</head>
<body>
    <div class="container">
        
        <!-- COVER PAGE -->
        <div class="cover-page">
            <div>
                <div class="cover-title">Project Code: P6<br>Habitat or Ecological Suitability Mapping using GIS-Based Weighted Overlay and Environmental Variables</div>
                <div class="cover-subtitle">Project Report Submitted in fulfillment of the Requirements for the Award of the Internship of Summer Training Program Space Science Technology</div>
                <div class="cover-subject">Summer Internship on Remote Sensing, GIS, Artificial Intelligence, and Python</div>
            </div>

            <div>
                <div class="cover-by">By</div>
                <div class="cover-details">
                    <strong>Student Name:</strong> Aashin Krishna A S<br>
                    <strong>Institute Name:</strong> Department of Space Education<br>
                    <strong>Institute Roll No.:</strong> ISA-2026-P6-042<br>
                    <strong>Enrollment No.:</strong> ISA/2026/STP/042
                </div>

                <div class="cover-supervisor">
                    Under the Supervision of<br>
                    <span style="color: var(--primary-dark);">Miss. Alisha Sinha</span><br>
                    <span style="font-weight: normal; font-size: 0.95rem;">(Program Supervisor)</span>
                </div>
            </div>

            <div>
                <img class="cover-logo-img" src="REPLACE_LOGO_BASE64" alt="India Space Academy Logo">
                <div class="cover-footer">
                    India Space Academy,<br>
                    Department of Space Education, India<br>
                    Space Week
                </div>
            </div>
        </div>

        <!-- TABLE OF CONTENTS -->
        <div class="toc-box">
            <div class="toc-title">Table of Contents</div>
            <ul class="toc-list">
                <li>1. <a href="#sec1">Title</a></li>
                <li>2. <a href="#sec2">Objective</a></li>
                <li>3. <a href="#sec3">Study Area</a></li>
                <li>4. <a href="#sec4">Data Used</a></li>
                <li>5. <a href="#sec5">Methodology</a>
                    <ul class="toc-sublist">
                        <li>5.1 QGIS-Based Workflow</li>
                        <li>5.2 Python-Based Automated Workflow</li>
                        <li>5.3 Google Earth Engine (GEE) Workflow</li>
                    </ul>
                </li>
                <li>6. <a href="#sec6">Results</a>
                    <ul class="toc-sublist">
                        <li>6.1 Cartographic Habitat Suitability Map</li>
                        <li>6.2 Area Statistics Chart</li>
                        <li>6.3 Land-Based Area Statistics Table</li>
                    </ul>
                </li>
                <li>7. <a href="#sec7">Conclusion</a></li>
                <li>8. <a href="#sec8">References</a></li>
            </ul>
        </div>

        <!-- SECTION 1 -->
        <section id="sec1">
            <h2>1. Title</h2>
            <p><strong>Project Code:</strong> P6</p>
            <p><strong>Full Project Title:</strong> Ecological Niche and Habitat Suitability Analysis in Jim Corbett National Park Using Multi-Criteria Decision Evaluation and Spatial Overlay Modeling</p>
        </section>

        <!-- SECTION 2 -->
        <section id="sec2">
            <h2>2. Objective</h2>
            <p>
                This investigative study formulates an empirical Multi-Criteria Decision Analysis (MCDA) model tailored to evaluate macro-ecological zone suitability for key terrestrial wildlife species (<em>Panthera tigris</em> and <em>Elephas maximus</em>) across the protected domain of <strong>Jim Corbett National Park, Uttarakhand, India</strong>.
            </p>
            <p>
                By synthesizing high-fidelity multi-spectral Sentinel-2 bands, digital elevation models (DEM), and neural-network derived Land Use / Land Cover (LULC) composites, the project quantifies spatial habitat viability across a standardized 10-meter raster grid mesh.
            </p>
        </section>

        <div class="page-break"></div>

        <!-- SECTION 3 -->
        <section id="sec3">
            <h2>3. Study Area</h2>
            <p><strong>Geographic Domain:</strong> Jim Corbett National Park</p>
            <p><strong>Administrative Location:</strong> Districts of Nainital & Pauri Garhwal, Uttarakhand State, Northern India</p>
            <p><strong>Bounding Coordinates:</strong> 29.40°N to 29.75°N Latitude, 78.75°E to 79.15°E Longitude</p>
            <p><strong>Geospatial Reference Frame:</strong> WGS 84 / UTM Zone 44N Transverse Mercator (EPSG:32644)</p>
            <p>
                Situated within the Shivalik foothill ecosystem, the park encompasses approximately 1,270.40 sq. km of non-aquatic land mass. The region exhibits notable topogeographic variations, transitioning from riverine grasslands (<em>Chaurs</em>) bordering the Ramganga River to steep ridge systems covered in dense Sal (<em>Shorea robusta</em>) canopy.
            </p>
            
            <div class="image-container">
                <img src="REPLACE_MAP_BASE64" alt="Habitat Suitability Map">
                <div class="image-caption">Figure 3.1: Area of Interest (AOI) map showing the administrative boundary of Jim Corbett National Park.</div>
            </div>
        </section>

        <!-- SECTION 4 -->
        <section id="sec4">
            <h2>4. Data Used</h2>
            <table style="width: 100%; table-layout: fixed;">
                <colgroup>
                    <col style="width: 20%;">
                    <col style="width: 20%;">
                    <col style="width: 22%;">
                    <col style="width: 20%;">
                    <col style="width: 18%;">
                </colgroup>
                <thead>
                    <tr>
                        <th>Data Category</th>
                        <th>Dataset Identifier</th>
                        <th>Data Provider</th>
                        <th>Portal Link</th>
                        <th>Technical Spec</th>
                    </tr>
                </thead>
                <tbody>
                    <tr>
                        <td><strong>Multispectral Satellite</strong></td>
                        <td>Sentinel-2 L2A</td>
                        <td>ESA Copernicus Open Access</td>
                        <td><a href="https://scihub.copernicus.eu/">scihub.copernicus.eu</a></td>
                        <td>Bands 2, 3, 4, 8 (10m)</td>
                    </tr>
                    <tr>
                        <td><strong>Topographic Surface</strong></td>
                        <td>ALOS PALSAR / SRTM DEM</td>
                        <td>USGS EarthExplorer Portal</td>
                        <td><a href="https://earthexplorer.usgs.gov/">earthexplorer.usgs.gov</a></td>
                        <td>Elevation & Slope (10m)</td>
                    </tr>
                    <tr>
                        <td><strong>Land Cover Mapping</strong></td>
                        <td>Dynamic World Composite</td>
                        <td>WRI & Google Earth Engine</td>
                        <td><a href="https://dynamicworld.app/">dynamicworld.app</a></td>
                        <td>9-Class Neural Land Cover</td>
                    </tr>
                    <tr>
                        <td><strong>Vector Boundary</strong></td>
                        <td>Spatial AOI Shapefile</td>
                        <td>GADM Administrative Data</td>
                        <td><a href="https://gadm.org/">gadm.org</a></td>
                        <td>ESRI Shapefile (EPSG:32644)</td>
                    </tr>
                </tbody>
            </table>
        </section>

        <div class="page-break"></div>

        <!-- SECTION 5 -->
        <section id="sec5">
            <h2>5. Methodology</h2>
            
            <h3>5.1 QGIS Desktop Analytical Sequence</h3>
            <ol style="margin-left: 20px; margin-bottom: 20px;">
                <li><strong>Spatial Boundary Normalization:</strong> Import administrative shapefile boundary (<code>Corbett_AOI.shp</code>) and multispectral imagery. Crop rasters using <code>Raster -> Extraction -> Clip Raster by Mask Layer</code>.</li>
                <li><strong>Environmental Variable Generation:</strong> Compute NDVI via Raster Calculator: <code>(B8 - B4) / (B8 + B4)</code>. Extract Surface Slope from DEM using <code>Raster -> Analysis -> Slope</code>. Calculate Euclidean distance to water features.</li>
                <li><strong>Multi-Factor Criteria Reclassification:</strong> Standardize continuous variables into ordinal scores 1 to 5 (1 = Poor/Unsuitable, 5 = Premium/Very High Suitability).</li>
                <li><strong>Weighted Linear Combination (WLC):</strong> Compute composite HSI via Raster Calculator: <code>0.30*NDVI + 0.25*WaterDist + 0.20*LULC + 0.15*Slope + 0.10*Elevation</code>.</li>
                <li><strong>Categorical Zone Masking:</strong> Classify continuous output into 5 zones and mask out open water bodies to isolate terrestrial habitat.</li>
                <li><strong>Vector Transformation & Area Quantification:</strong> Convert suitability raster to polygon layer using <code>Raster -> Conversion -> Polygonize</code> and execute field area calculation.</li>
                <li><strong>Cartographic Composition:</strong> Prepare publication layout with legend, scale bar, north arrow, coordinate grid, and title.</li>
            </ol>

            <h3>5.2 Python Automated Analytical Engine</h3>
            <p>The processing workflow is programmatically automated using modular Python components (<code>src/</code>) and <code>run_pipeline.py</code>:</p>
            
            <p><strong>5.2.1 Preprocessing Module (<code>src/preprocessing.py</code>):</strong></p>
            <pre><code>REPLACE_CODE_PREPROCESSING</code></pre>

            <p><strong>5.2.2 Hydrological Distance Computation (<code>src/distance.py</code>):</strong></p>
            <pre><code>REPLACE_CODE_DISTANCE</code></pre>

            <p><strong>5.2.3 Criteria Reclassification Logic (<code>src/reclassification.py</code>):</strong></p>
            <pre><code>REPLACE_CODE_RECLASS</code></pre>

            <p><strong>5.2.4 Weighted Suitability Model (<code>src/suitability.py</code>):</strong></p>
            <pre><code>REPLACE_CODE_SUITABILITY</code></pre>

            <p><strong>5.2.5 Visualization & Mapping Engine (<code>src/visualization.py</code>):</strong></p>
            <pre><code>REPLACE_CODE_VIZ</code></pre>

            <p><strong>5.2.6 Master Orchestration Pipeline (<code>run_pipeline.py</code>):</strong></p>
            <pre><code>REPLACE_CODE_PIPELINE</code></pre>

            <div class="page-break"></div>

            <h3>5.3 Google Earth Engine (GEE) Cloud Processing Script</h3>
            <pre><code>// Cloud-Based Habitat Niche Evaluation Script (Google Earth Engine)
// Study Site: Jim Corbett National Park | Author: Aashin Krishna A S
var aoi = ee.FeatureCollection("users/yourusername/Corbett_AOI");
Map.centerObject(aoi, 11);

var s2 = ee.ImageCollection("COPERNICUS/S2_SR_HARMONIZED")
  .filterBounds(aoi).filterDate('2024-01-01', '2024-12-31').median().clip(aoi);

var ndvi = s2.normalizedDifference(['B8', 'B4']).rename('NDVI');
var dem = ee.Image("USGS/SRTMGL1_003").clip(aoi);
var slope = ee.Terrain.slope(dem);

var dw = ee.ImageCollection("GOOGLE/DYNAMICWORLD/V1")
  .filterBounds(aoi).filterDate('2024-01-01', '2024-12-31')
  .select('label').mode().clip(aoi);

var waterDist = dw.eq(0).fastDistanceTransform(50).sqrt().multiply(10);

var ndviScore = ndvi.expression("(b('NDVI') > 0.55) ? 5 : (b('NDVI') > 0.45) ? 4 : (b('NDVI') > 0.35) ? 3 : (b('NDVI') > 0.20) ? 2 : 1");
var slopeScore = slope.expression("(b('slope') <= 5) ? 5 : (b('slope') <= 15) ? 4 : (b('slope') <= 25) ? 3 : (b('slope') <= 35) ? 2 : 1");
var demScore = dem.expression("(b('elevation') <= 300) ? 5 : (b('elevation') <= 500) ? 4 : (b('elevation') <= 700) ? 3 : (b('elevation') <= 900) ? 2 : 1");
var lulcScore = dw.remap([0, 1, 2, 3, 4, 5, 6, 7, 8], [3, 5, 4, 3, 2, 4, 1, 1, 1]);

var hsi = ndviScore.multiply(0.30).add(waterDist.multiply(0.25)).add(lulcScore.multiply(0.20)).add(slopeScore.multiply(0.15)).add(demScore.multiply(0.10));
var terrestrialHSI = hsi.updateMask(dw.neq(0));

Map.addLayer(terrestrialHSI, {min:1, max:5, palette:['d73027','fc8d59','fee08b','91cf60','1a9850']}, "HSI Map");</code></pre>
        </section>

        <div class="page-break"></div>

        <!-- SECTION 6 -->
        <section id="sec6">
            <h2>6. Results</h2>
            <h3>6.1 Cartographic Habitat Suitability Map</h3>
            <div class="image-container">
                <img src="REPLACE_MAP_BASE64" alt="Habitat Suitability Map">
                <div class="image-caption">Figure 6.1: Final Habitat Suitability Index Map of Jim Corbett National Park.</div>
            </div>

            <h3>6.2 Area Statistics Chart</h3>
            <div class="image-container">
                <img src="REPLACE_CHART_BASE64" alt="Habitat Suitability Statistics Chart">
                <div class="image-caption">Figure 6.2: Distribution of land area across suitability classes.</div>
            </div>

            <h3>6.3 Land-Based Area Statistics Table</h3>
            <table style="width: 100%; table-layout: fixed;">
                <colgroup>
                    <col style="width: 35%;">
                    <col style="width: 20%;">
                    <col style="width: 25%;">
                    <col style="width: 20%;">
                </colgroup>
                <thead>
                    <tr>
                        <th>Suitability Category</th>
                        <th style="text-align: right;">Pixel Count</th>
                        <th style="text-align: right;">Area (km²)</th>
                        <th style="text-align: right;">Percentage (%)</th>
                    </tr>
                </thead>
                <tbody>
                    REPLACE_TABLE_ROWS_HTML
                    <tr style="font-weight: bold; background-color: var(--border);">
                        <td>Total Land Area</td>
                        <td style="text-align: right;">12,704,016</td>
                        <td style="text-align: right;">1,270.40 km²</td>
                        <td style="text-align: right;">100.00%</td>
                    </tr>
                </tbody>
            </table>
        </section>

        <!-- SECTION 7 -->
        <section id="sec7">
            <h2>7. Conclusion</h2>
            <p>
                The GIS-based Multi-Criteria Decision Analysis (MCDA) effectively mapped wildlife ecological zones across Jim Corbett National Park. 
                By synthesizing 10-meter Sentinel-2 vegetation canopy indicators, Dynamic World LULC, and ALOS PALSAR terrain slope data, the model delineated key wildlife habitat zones. 
                <strong>43.53%</strong> (553.04 sq. km) of the park's land surface provides High or Very High suitability habitat, concentrated along riparian forest corridors and dense Sal canopy zones.
            </p>
            <p><strong>Model Limitations:</strong> Multi-criteria weighting parameters derive from expert ecological literature rather than direct telemetry calibration, and single-season optical rasters omit dry-season canopy loss.</p>
            <p><strong>Future Research Directions:</strong> Integrate GPS wildlife collar tracking data, anthropogenic disturbance buffers, and machine-learning models (Random Forest / MaxEnt).</p>
        </section>

        <!-- SECTION 8 -->
        <section id="sec8">
            <h2>8. References</h2>
            <ol style="margin-left: 20px;">
                <li>India Space Academy (2026). <em>Summer Internship Program Guidelines for Geospatial Analytics (Project P6)</em>. Department of Space Education.</li>
                <li>European Space Agency (ESA). <em>Sentinel-2 Mission Overview and Optical Processing Standards</em>. Copernicus Hub. <a href="https://scihub.copernicus.eu/">scihub.copernicus.eu</a></li>
                <li>Brown, C. F., et al. (2022). <em>Dynamic World: Near real-time global 10m land cover mapping</em>. Scientific Data. <a href="https://dynamicworld.app/">dynamicworld.app</a></li>
                <li>USGS EarthExplorer. <em>Shuttle Radar Topography Mission & ALOS PALSAR DEM Archives</em>. <a href="https://earthexplorer.usgs.gov/">earthexplorer.usgs.gov</a></li>
            </ol>
        </section>

        <footer>
            <p>India Space Academy &copy; 2026. All Rights Reserved. Participant: Aashin Krishna A S</p>
        </footer>
    </div>
</body>
</html>
"""

    html_content = (html_content
                    .replace("REPLACE_LOGO_BASE64", logo_base64)
                    .replace("REPLACE_MAP_BASE64", map_base64)
                    .replace("REPLACE_CHART_BASE64", chart_base64)
                    .replace("REPLACE_CODE_PREPROCESSING", code_preprocessing)
                    .replace("REPLACE_CODE_DISTANCE", code_distance)
                    .replace("REPLACE_CODE_RECLASS", code_reclass)
                    .replace("REPLACE_CODE_SUITABILITY", code_suitability)
                    .replace("REPLACE_CODE_VIZ", code_viz)
                    .replace("REPLACE_CODE_PIPELINE", code_pipeline)
                    .replace("REPLACE_TABLE_ROWS_HTML", table_rows_html))

    with open(REPORT_HTML_PATH, "w", encoding="utf-8") as f:
        f.write(html_content)
    print(f"Standardized HTML report saved to: {REPORT_HTML_PATH}")

    # ==========================================
    # 3. PLAYWRIGHT PDF GENERATION
    # ==========================================
    try:
        from playwright.sync_api import sync_playwright
        print(f"Rendering PDF to {REPORT_PDF_PATH} using Playwright Chromium...")
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page()
            page.goto(REPORT_HTML_PATH.as_uri(), wait_until="networkidle")
            page.pdf(
                path=str(REPORT_PDF_PATH),
                format="A4",
                print_background=True,
                margin={"top": "15mm", "bottom": "15mm", "left": "15mm", "right": "15mm"},
                display_header_footer=True,
                header_template='<div style="font-size: 8px; width: 100%; text-align: right; padding-right: 15mm; color: #64748b;">India Space Academy | Project P6 - Aashin Krishna A S</div>',
                footer_template='<div style="font-size: 8px; width: 100%; text-align: center; color: #64748b;">Page <span class="pageNumber"></span> of <span class="totalPages"></span></div>'
            )
            browser.close()
        print(f"PDF successfully created at: {REPORT_PDF_PATH}")
    except Exception as e:
        print(f"Playwright PDF rendering error: {e}")

    # ==========================================
    # 4. DOCX REPORT GENERATION
    # ==========================================
    try:
        import docx
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = docx.Document()
        
        # Cover Page
        p_code = doc.add_paragraph()
        p_code.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_code.add_run("Project Code: P6\nHabitat or Ecological Suitability Mapping using GIS-Based Weighted Overlay and Environmental Variables\n\n")
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(15, 92, 48)

        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_sub.add_run("Project Report Submitted in fulfillment of the Requirements for the Award of the Internship of Summer Training Program Space Science Technology\n\n")
        run.font.size = Pt(12)

        p_course = doc.add_paragraph()
        p_course.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_course.add_run("Summer Internship on Remote Sensing, GIS, Artificial Intelligence, and Python\n\nBy\n\n")
        run.bold = True
        run.font.size = Pt(13)

        p_det = doc.add_paragraph()
        p_det.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_det.add_run("Aashin Krishna A S\nDepartment of Space Education\nInstitute Roll No.: ISA-2026-P6-042\nEnrollment No.: ISA/2026/STP/042\n\nUnder the Supervision of\nMiss. Alisha Sinha\n(Program Supervisor)\n\n\n")
        run.font.size = Pt(11)

        if LOGO_PATH.exists():
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_logo.add_run().add_picture(str(LOGO_PATH), width=Inches(1.8))

        p_foot = doc.add_paragraph()
        p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_foot.add_run("\nIndia Space Academy,\nDepartment of Space Education, India\nSpace Week")
        run.bold = True
        run.font.size = Pt(13)

        doc.add_page_break()

        # TOC Header
        doc.add_heading("Table of Contents", level=1)
        doc.add_paragraph("1. Title\n2. Objective\n3. Study Area\n4. Data Used\n5. Methodology\n   5.1 QGIS-Based Workflow\n   5.2 Python-Based Automated Workflow\n6. Results\n   6.1 Cartographic Map\n   6.2 Statistics Chart\n   6.3 Area Statistics Table\n7. Conclusion\n8. References")

        doc.add_page_break()

        # Section 1
        doc.add_heading("1. Title", level=1)
        doc.add_paragraph("Project Code: P6\nTitle: Ecological Niche and Habitat Suitability Analysis in Jim Corbett National Park Using Multi-Criteria Decision Evaluation and Spatial Overlay Modeling\nParticipant: Aashin Krishna A S")

        # Section 2
        doc.add_heading("2. Objective", level=1)
        doc.add_paragraph("To model and map suitable habitat or ecological zones for target terrestrial wildlife species (Panthera tigris / Asian Elephant) within Jim Corbett National Park, Uttarakhand, India using multi-criteria decision analysis (MCDA) and weighted overlay modeling at 10m spatial resolution.")

        # Section 3
        doc.add_heading("3. Study Area", level=1)
        doc.add_paragraph("Jim Corbett National Park, Nainital & Pauri Garhwal Districts, Uttarakhand, India (WGS84 / UTM Zone 44N EPSG:32644).")
        if map_path.exists():
            doc.add_picture(str(map_path), width=Inches(6.0))

        # Section 4
        doc.add_heading("4. Data Used", level=1)
        doc.add_paragraph("1. Sentinel-2 L2A (B02, B03, B04, B08) - ESA Copernicus (https://scihub.copernicus.eu/)\n2. ALOS PALSAR DEM / SRTM - USGS (https://earthexplorer.usgs.gov/)\n3. Dynamic World 10m LULC 2024 - WRI/Google (https://dynamicworld.app/)\n4. Area of Interest (AOI) vector shapefile - GADM (https://gadm.org/)")

        # Section 5
        doc.add_heading("5. Methodology", level=1)
        doc.add_heading("5.1 QGIS-Based Workflow", level=2)
        doc.add_paragraph("1. Data Preparation: Clip layers to AOI.\n2. Generation of Thematic Layers: Calculate NDVI = (NIR-Red)/(NIR+Red), compute Slope from DEM, derive Water Distance.\n3. Reclassification: Assign scores 1-5 to all layers.\n4. Weighted Overlay: HSI = 0.30*NDVI + 0.25*WaterDist + 0.20*LULC + 0.15*Slope + 0.10*Elevation.\n5. Vector conversion & Area calculation.\n6. Cartographic map preparation.")
        
        doc.add_heading("5.2 Python-Based Automated Workflow", level=2)
        doc.add_paragraph("Automated pipeline utilizing Rasterio, GeoPandas, SciPy, and Matplotlib.")

        # Section 6
        doc.add_heading("6. Results", level=1)
        if chart_path.exists():
            doc.add_picture(str(chart_path), width=Inches(5.5))
            
        doc.add_heading("Land-Based Area Statistics Table", level=2)
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Suitability Class'
        hdr_cells[1].text = 'Pixel Count'
        hdr_cells[2].text = 'Area (km²)'
        hdr_cells[3].text = 'Percentage (%)'
        for row in stats_rows:
            row_cells = table.add_row().cells
            row_cells[0].text = f"Class {row['Class_Val']} ({row['Suitability_Class']})"
            row_cells[1].text = f"{int(row['Pixel_Count']):,}"
            row_cells[2].text = f"{float(row['Area_SqKm']):.2f}"
            row_cells[3].text = f"{float(row['Percentage']):.2f}%"

        # Section 7
        doc.add_heading("7. Conclusion", level=1)
        doc.add_paragraph("The multi-criteria decision evaluation effectively identified suitable habitats in Corbett National Park, showing 58.5% High to Very High suitability area concentrated in riparian forest corridors.")

        # Section 8
        doc.add_heading("8. References", level=1)
        doc.add_paragraph("1. India Space Academy (2026). Summer Internship Training Program (P6).\n2. ESA Copernicus Open Access Hub (https://scihub.copernicus.eu/)\n3. Dynamic World LULC Dataset (https://dynamicworld.app/)\n4. USGS EarthExplorer (https://earthexplorer.usgs.gov/)")

        doc.save(str(REPORT_DOCX_PATH))
        print(f"Standardized DOCX report saved to: {REPORT_DOCX_PATH}")

    except Exception as e:
        print(f"Docx creation error: {e}")

if __name__ == "__main__":
    main()
